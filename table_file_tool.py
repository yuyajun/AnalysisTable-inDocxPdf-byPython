#!/usr/bin/python
# -*- coding: UTF-8 -*-
import os
import re
from docx import Document
import pdfplumber
import pandas as pd


# 定义需统计的相关信息
columns=["应聘岗位", "姓名", "出生年月", "性别", "联系方式", "工作单位"]
df = pd.DataFrame(columns=columns)


# 所需处理文件存储的路径（文件夹路径）
path = "/Users/yuyajun/Desktop/test_file"

# 计数器统计处理文档个数
count = 0
# 遍历path下所有文档
for dirpath, dirnames, filenames in os.walk(path):
    for filepath in filenames:
        # 当前处理文件路径
        count = count + 1
        doc_path = os.path.join(dirpath, filepath)
        print("\n")
        print(count, "--------现在处理的文件是: ", doc_path)

        # 分析docx后缀的文件
        if doc_path[-4:] == "docx":
            document = Document(doc_path)

            # 从第一页第二行提取应聘岗位
            job_for = document.paragraphs[1].text
            job_for = re.split(':|：', job_for)[-1].replace(" ", "")

            # 获取word里面表格及其数量
            tables = document.tables
            tables_len = len(tables)

            # 该项目场景下只允许有一个表格
            if tables_len == 1:
                # 获取word中一级表格
                table = tables[0]

                # table.cell(i, j) 获取对应信息
                # replace(" ","").replace("\n","") 删除字符串中所有空格和换行符
                name = table.cell(0, 1).text.replace(" ","").replace("\n","")
                birth = table.cell(0, 6).text.replace(" ","").replace("\n","")
                gender = table.cell(0, 3).text.replace(" ","").replace("\n","")
                cellphone = table.cell(5, 1).text.replace(" ","").replace("\n","")

                # 获取所有工作单位
                tag = -1
                temp_data_work = []
                for row in table.rows:
                    for cell in row.cells:
                        # 内容不为空
                        if cell.text:
                            data = cell.text.replace(" ","").replace("\n","")
                            if data == "备注":
                                tag = 0
                            elif data.startswith("奖惩情况"):
                                tag = 1
                            # 将"备注"和"奖惩情况"中间的行(即工作经历)存储
                            elif tag == 0:
                                temp_data_work.append(data)

                # 根据表格特点，获取所有工作时间+工作单位
                print(temp_data_work)
                work_unit_list = []
                for i in range(len(temp_data_work)):
                    if i%9 == 0:
                        work_unit_list.append(temp_data_work[i] + ": " + temp_data_work[i+1])
                work_unit = ", ".join(work_unit_list)
                print(work_unit)

                # 内容不能为空
                if not (job_for and name and birth and gender and cellphone and work_unit):
                    print(count, "take care 表格格式/填写有误:", doc_path)
                else:
                    # 插入数据
                    data = [job_for, name, birth, gender, cellphone, work_unit]
                    df_add = pd.DataFrame([data], columns=columns)
                    df = pd.concat([df, df_add], ignore_index=True)
                    print(count, "Good job! -- docx")
            else:
                print(count, "warning: 文件中表格个数不正确, 个数为:", tables_len, ",文件名:", doc_path)

        # 分析pdf后缀的文件
        elif doc_path[-3:] == "pdf":
            pdf = pdfplumber.open(doc_path)

            # 本场景中会重复出现"姓名""出生年月"，用tag标来防止内容被覆盖。
            tag = -1

            # 临时存储所有工作经历
            temp_data_work = []
            for i in range(len(pdf.pages)):
                page = pdf.pages[i]
                # 从第一页第二行提取应聘岗位
                if i == 0:
                    job_for = page.extract_text().split("\n")[1]
                    job_for = re.split(':|：', job_for)[-1].replace(" ","")
                # 获取当前页面表格
                table = page.extract_tables()

                for j in range(len(table)):
                    for k in range(len(table[j])):
                       for m in range(len(table[j][k])):
                           # 获取第j个表格的第k行第m列（从0开始计数）的内容
                            data = table[j][k][m]
                            if data:
                                # replace(" ","").replace("\n","") 删除字符串中所有空格和换行符
                                data = data.replace(" ","").replace("\n","")
                                if data == "姓名" and tag == -1:
                                    name = table[j][k][m+1]
                                elif data == "出生年月" and tag == -1:
                                    birth = table[j][k][m+1]
                                elif data == "性别":
                                    gender = table[j][k][m+1]
                                elif data == "手机号码":
                                    cellphone = table[j][k][m+1]
                                elif data == "备注":
                                    tag = 1
                                elif data == "奖惩情况（荣誉）" or data == "奖惩情况":
                                    tag = 2
                       # 将"备注"和"奖惩情况"中间的行(即工作经历)存储
                       if tag == 1:
                           temp_data_work.append(table[j][k])

            # 获取所有工作时间+工作单位
            work_unit_list = []
            length = len(temp_data_work)
            for i in range(length):
                # 第0行是['起止年月','所在学校/机构','所任职务','备注']
                if i != 0:
                    data = temp_data_work[i][0] + ": " + temp_data_work[i][1]
                    if data != ": ":
                        work_unit_list.append(data)
            work_unit = ", ".join(work_unit_list)

            # 内容不能为空
            if not (job_for and name and birth and gender and cellphone and work_unit):
                print(count, "take care 表格格式/填写有误:", doc_path)
            else:
                # 插入数据
                info = [job_for, name, birth, gender, cellphone, work_unit]
                df_add = pd.DataFrame([info], columns=columns)
                df = pd.concat([df, df_add], ignore_index=True)
                print(count, "Good job! -- pdf")

        # 需注意doc后缀文档，需转换为docx/pdf格式再进一步处理。
        elif doc_path[-3:] == "doc":
            print(count, "take care DOC文档:", doc_path)

        # 除doc/docx/pdf以外的文档都属于不符合规定文档，均忽略不处理。
        else:
            print(count, "warning: 文件后缀不合格: ", doc_path)

# 去重后，存储数据到汇总表
df.drop_duplicates(subset=["姓名", "联系方式"], inplace=True, ignore_index=True)
df.to_excel("/Users/yuyajun/Desktop/test_file/信息汇总表.xlsx")
